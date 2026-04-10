import sys
from docx import Document

def analyze_co_doc(path):
    doc = Document(path)
    print(f"Analyzing {path}...")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    # Extract headers and first few lines of each section
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') or p.text.isupper() and len(p.text) < 100:
            print(f"\n[HEADER] {p.text.strip()}")
        elif len(p.text.strip()) > 100:
            # Check for unit keywords
            if 'UNIT' in p.text.upper()[:20]:
                print(f"\n[UNIT DETECTED] {p.text.strip()[:100]}...")

if __name__ == "__main__":
    analyze_co_doc(r"resources\Theory\BCA Computer Organization Exam Prep.docx")
